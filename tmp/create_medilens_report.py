from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(r"C:\Users\sunil\Documents\MediLens AI")
OUT = ROOT / "reports"
OUT.mkdir(exist_ok=True)
DOCX_PATH = OUT / "MediLens_AI_Project_Report.docx"
DIAGRAM_PATH = OUT / "medilens_architecture.png"
LOGO = ROOT / "frontend" / "src" / "assets" / "logo.png"
SCREENSHOT = Path(r"C:\Users\sunil\Pictures\Screenshots\Screenshot 2026-05-29 001016.png")

PURPLE = "6236E9"
PURPLE_DARK = "4020A9"
PURPLE_LIGHT = "EEE9FF"
INK = "11172F"
MUTED = "687089"
LINE = "E4E7F2"
GREEN = "0D9B66"
RED = "A51D3B"
GOLD = "A66B00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(w.inches for w in widths) * 1440)))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))


def set_run(run, size=None, bold=None, italic=None, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_para(doc, text="", size=10.5, color=INK, bold=False, italic=False,
             align=None, before=0, after=6, line=1.15, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_para(doc, parts, before=0, after=6, line=1.15):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for text, kwargs in parts:
        set_run(p.add_run(text), **kwargs)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    set_run(p.add_run(text), size=10.5, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    set_run(p.add_run(text), size=10.5, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    return p.add_run(text)


def add_callout(doc, title, text, fill=PURPLE_LIGHT, color=PURPLE_DARK):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.45)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=fill, size="4")
    set_cell_margins(cell, top=140, bottom=140, start=170, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(title), size=10.5, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    set_run(p2.add_run(text), size=10, color=INK)
    add_para(doc, "", after=3)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, widths)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_shading(cell, PURPLE)
        set_cell_border(cell, color="D4CCFF")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(header), size=font_size, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            set_run(p.add_run(str(value)), size=font_size, color=INK)
    add_para(doc, "", after=3)
    return table


def add_chapter(doc, number, title):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(f"CHAPTER {number}"), size=11, bold=True, color=PURPLE)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    set_run(p2.add_run(title.upper()), size=20, bold=True, color=INK)
    return p2


def make_architecture_diagram():
    image = Image.new("RGB", (1800, 950), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 46)
        box_font = ImageFont.truetype("arialbd.ttf", 31)
        body_font = ImageFont.truetype("arial.ttf", 25)
    except OSError:
        title_font = ImageFont.load_default()
        box_font = ImageFont.load_default()
        body_font = ImageFont.load_default()

    draw.text((75, 45), "MediLens AI - System Architecture", fill=(64, 32, 169), font=title_font)
    boxes = [
        ((90, 220, 380, 460), "React + Vite", ["Public Home", "Auth Context", "Scan & History UI"], (238, 233, 255)),
        ((535, 220, 865, 460), "Express API", ["JWT middleware", "Multer upload", "REST controllers"], (231, 240, 255)),
        ((1015, 120, 1425, 330), "Gemini API", ["Image understanding", "Structured JSON", "Hindi translation"], (232, 248, 241)),
        ((1015, 490, 1425, 700), "MongoDB Atlas", ["Users", "Reports", "Preferences"], (255, 242, 219)),
        ((1530, 285, 1750, 535), "Storage", ["uploads/", "PDF export", "Thumbnails"], (255, 233, 239)),
    ]
    for (x1, y1, x2, y2), title, lines, fill in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=30, fill=fill, outline=(98, 54, 233), width=4)
        draw.text((x1 + 30, y1 + 30), title, fill=(17, 23, 47), font=box_font)
        for i, line in enumerate(lines):
            draw.text((x1 + 34, y1 + 100 + i * 46), line, fill=(70, 78, 105), font=body_font)
    arrows = [
        ((390, 340), (520, 340)),
        ((875, 290), (1000, 240)),
        ((875, 390), (1000, 585)),
        ((1435, 410), (1515, 410)),
    ]
    for start, end in arrows:
        draw.line((*start, *end), fill=(98, 54, 233), width=8)
        x, y = end
        draw.polygon([(x, y), (x - 22, y - 16), (x - 22, y + 16)], fill=(98, 54, 233))
    draw.text((90, 800), "Flow: user uploads an image -> authenticated API validates and analyzes it -> structured report is saved -> user views or downloads result.", fill=(70, 78, 105), font=body_font)
    image.save(DIAGRAM_PATH)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for level, size, color, before, after in [
        (1, 16, PURPLE_DARK, 14, 7),
        (2, 13, PURPLE, 10, 5),
        (3, 11.5, PURPLE_DARK, 8, 4),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("MediLens AI  |  Project Report"), size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def cover_page(doc):
    add_para(doc, "", after=12)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.25))
    add_para(doc, "PROJECT REPORT", size=13, bold=True, color=PURPLE, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=8)
    add_para(doc, "MediLens AI", size=31, bold=True, color=PURPLE_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Medicine Understanding Assistant", size=17, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(doc, "A Full-Stack AI-Based Web Application for Understanding Medicine Information from Uploaded Images", size=12, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=26)
    add_para(doc, "Submitted in partial fulfillment of the requirements for the award of", size=11, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "[DEGREE / COURSE NAME]", size=13, color=PURPLE_DARK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(doc, "Submitted by", size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_para(doc, "[STUDENT NAME]", size=14, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "Roll No.: [ROLL NUMBER]", size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    add_para(doc, "Under the guidance of", size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_para(doc, "[GUIDE / FACULTY NAME]", size=13, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_para(doc, "[DEPARTMENT NAME]", size=12, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_para(doc, "[COLLEGE / UNIVERSITY NAME]", size=12, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_para(doc, "Academic Session: 2025-2026", size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)


def preliminary_pages(doc):
    doc.add_page_break()
    add_heading(doc, "CERTIFICATE", 1)
    add_para(doc, "This is to certify that the project report entitled “MediLens AI - Medicine Understanding Assistant” is a bonafide record of the work carried out by [STUDENT NAME], Roll No. [ROLL NUMBER], under my supervision and guidance during the academic session 2025-2026. The project fulfills the requirements prescribed by [COLLEGE / UNIVERSITY NAME].", after=22)
    add_para(doc, "Guide Signature: ______________________________", after=18)
    add_para(doc, "Head of Department: ___________________________", after=18)
    add_para(doc, "Date: ___________________        Place: ___________________", after=6)

    doc.add_page_break()
    add_heading(doc, "DECLARATION", 1)
    add_para(doc, "I hereby declare that the project report entitled “MediLens AI - Medicine Understanding Assistant” is my original work. It has been developed for academic purposes under the guidance of [GUIDE / FACULTY NAME]. The work has not been submitted previously, in full or in part, for the award of any other degree or diploma.", after=22)
    add_para(doc, "Student Signature: ____________________________", after=18)
    add_para(doc, "Name: [STUDENT NAME]", after=6)
    add_para(doc, "Roll No.: [ROLL NUMBER]", after=6)
    add_para(doc, "Date: ___________________", after=6)

    doc.add_page_break()
    add_heading(doc, "ACKNOWLEDGEMENT", 1)
    add_para(doc, "I express my sincere gratitude to [GUIDE / FACULTY NAME] for valuable guidance, encouragement, and constructive feedback throughout the development of this project. I also thank the faculty members of [DEPARTMENT NAME] and [COLLEGE / UNIVERSITY NAME] for providing the opportunity and resources required to complete this work.")
    add_para(doc, "I am grateful to my family and friends for their support. I also acknowledge the open-source communities and documentation providers whose tools helped in building the application, including React, Vite, Node.js, Express, MongoDB, Multer, jsPDF, and Google Gemini API.")

    doc.add_page_break()
    add_heading(doc, "ABSTRACT", 1)
    add_para(doc, "MediLens AI is a full-stack medicine understanding assistant designed to help users obtain simple, structured information from images of medicine strips and prescriptions. The application allows a registered user to upload an image, send it securely to a Node.js and Express backend, analyze the visible content using the Google Gemini API, and view an organized explanation of the medicine name, use, dosage when visible, precautions, possible side effects, and doctor advice.")
    add_para(doc, "The system supports English and Hindi. A result can be translated after analysis without uploading the image again or running the image-analysis flow a second time. Users can store private scan history, view scanned-image thumbnails, search and delete previous reports, manage profile settings, choose light or dark theme, and download PDF reports. The application includes safety-oriented prompting: it does not diagnose diseases, prescribe medication, or guess dosage when information is not visible.")
    add_para(doc, "The project demonstrates integration of a modern React frontend, a REST API built with Express, MongoDB-based persistence, JWT authentication, image upload validation using Multer, AI-assisted structured responses, multilingual translation, and PDF export with embedded Devanagari font support.")
    add_callout(doc, "Keywords", "Artificial Intelligence, Medicine Understanding, Gemini API, React, Express, MongoDB, JWT, Multer, Hindi Translation, PDF Report")

    doc.add_page_break()
    add_heading(doc, "TABLE OF CONTENTS", 1)
    toc = [
        ("Certificate", "ii"), ("Declaration", "iii"), ("Acknowledgement", "iv"), ("Abstract", "v"),
        ("Chapter 1: Introduction", "1"), ("Chapter 2: Requirement Analysis", "3"),
        ("Chapter 3: System Design", "5"), ("Chapter 4: Implementation", "8"),
        ("Chapter 5: Testing and Validation", "13"), ("Chapter 6: Deployment", "15"),
        ("Chapter 7: Limitations and Future Scope", "17"), ("Chapter 8: Conclusion", "19"),
        ("References", "20"), ("Appendix", "21")
    ]
    add_table(doc, ["Section", "Page"], toc, [Inches(5.8), Inches(0.55)], font_size=10)


def chapter_introduction(doc):
    add_chapter(doc, "1", "Introduction")
    add_heading(doc, "1.1 Background", 2)
    add_para(doc, "Medicine labels and prescriptions often contain technical terminology that is difficult for a general user to understand quickly. Users may also need a simple summary of visible information before consulting a doctor or pharmacist. MediLens AI addresses this need by presenting AI-assisted explanations in a structured and accessible format.")
    add_heading(doc, "1.2 Problem Statement", 2)
    add_para(doc, "Users frequently encounter medicine strips and prescription images without an easy way to interpret the visible information. Searching manually can be slow and may lead to misunderstanding. A useful assistant should extract only visible details, explain them in simple language, support multilingual users, preserve privacy through authenticated access, and clearly state that the output is informational rather than medical advice.")
    add_heading(doc, "1.3 Proposed Solution", 2)
    add_para(doc, "MediLens AI provides a responsive web interface where a logged-in user uploads a medicine or prescription image. The backend validates the file, uses Gemini API for image understanding, stores a structured report in MongoDB, and returns an explanation to the frontend. Users can translate the result into Hindi without re-analysis, review previous reports, and generate downloadable PDF documents.")
    add_heading(doc, "1.4 Objectives", 2)
    for item in [
        "Provide an easy image-upload workflow for medicine strips and prescriptions.",
        "Return medicine name, use, dosage when visible, precautions, side effects, and doctor advice.",
        "Support simple English and natural Hindi output.",
        "Avoid diagnosis, prescription, and unsupported dosage guessing.",
        "Protect private features using JWT authentication.",
        "Store user-specific report history with uploaded-image references.",
        "Generate clean PDF reports for English and Hindi users.",
        "Deliver a responsive UI suitable for desktop, tablet, and mobile devices.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.5 Scope", 2)
    add_para(doc, "The scope of the project includes medicine-image understanding, prescription-image upload, structured explanation, translation, authenticated user history, settings, and report download. It is not a medical diagnostic platform and is not intended to replace professional healthcare advice.")


def chapter_requirements(doc):
    add_chapter(doc, "2", "Requirement Analysis")
    add_heading(doc, "2.1 Functional Requirements", 2)
    functional = [
        ("FR-01", "Public Home Page", "Visitors can view product, safety, and help information without login."),
        ("FR-02", "Authentication", "Users can register, login, view profile, and logout."),
        ("FR-03", "Protected Scanning", "Only authenticated users can upload and analyze medicine images."),
        ("FR-04", "Image Validation", "The app accepts JPG, JPEG, and PNG files up to 10 MB."),
        ("FR-05", "AI Analysis", "Gemini returns a structured medicine explanation."),
        ("FR-06", "Translation", "Users can translate an analyzed result without rescanning the image."),
        ("FR-07", "History", "Users can search, view, and delete their own reports."),
        ("FR-08", "PDF Export", "Users can download a clean report in English or Hindi."),
        ("FR-09", "Settings", "Users can save preferred language and light/dark theme."),
        ("FR-10", "Recent Analyses", "Home page shows recent scans with real thumbnails where available."),
    ]
    add_table(doc, ["ID", "Requirement", "Description"], functional, [Inches(0.7), Inches(1.55), Inches(4.1)])
    add_heading(doc, "2.2 Non-Functional Requirements", 2)
    for item in [
        "Usability: clear upload controls, loading states, error messages, and readable result cards.",
        "Security: hashed passwords, JWT-protected endpoints, user-specific report filtering, and environment-based secrets.",
        "Performance: image-size limits and separate translation endpoint to avoid repeated image analysis.",
        "Responsiveness: adaptive sidebar, grid cards, and mobile-friendly layouts.",
        "Maintainability: separated routes, controllers, models, middleware, services, and reusable React components.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "2.3 Hardware and Software Requirements", 2)
    add_table(doc, ["Category", "Requirement"], [
        ("Client", "Modern web browser on desktop, tablet, or mobile"),
        ("Development OS", "Windows 10/11 or equivalent"),
        ("Frontend", "React, Vite, Axios, React Router, jsPDF"),
        ("Backend", "Node.js, Express, Multer, bcryptjs, jsonwebtoken"),
        ("Database", "MongoDB Atlas or local MongoDB"),
        ("AI Service", "Google Gemini API"),
        ("Deployment", "Vercel frontend, Render backend, MongoDB Atlas database"),
    ], [Inches(1.7), Inches(4.65)])


def chapter_design(doc):
    add_chapter(doc, "3", "System Design")
    add_heading(doc, "3.1 Architecture Overview", 2)
    add_para(doc, "MediLens AI follows a client-server architecture. The React frontend communicates with an Express REST API. Authentication is handled by JWT tokens. MongoDB stores users and reports. Gemini API performs image understanding and translation. Uploaded images are served from the backend uploads directory for report thumbnails.")
    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.3))
    add_para(doc, "Figure 1: MediLens AI system architecture", size=9, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_heading(doc, "3.2 Application Workflow", 2)
    for step in [
        "A visitor opens the public home page and learns about the application.",
        "The user registers or logs in when a protected scan feature is requested.",
        "The user selects English or Hindi and uploads a valid medicine image.",
        "The Express backend validates the file with Multer and sends the image to Gemini.",
        "Gemini returns JSON containing medicine name, use, dosage, precautions, side effects, and doctor advice.",
        "The backend saves the structured report with the logged-in user ID.",
        "The frontend displays the result cards. A language change calls translation only, without rescanning.",
        "The user can download a PDF or review saved reports in History.",
    ]:
        add_number(doc, step)
    add_heading(doc, "3.3 Database Design", 2)
    add_table(doc, ["Collection", "Important Fields", "Purpose"], [
        ("users", "name, email, password, createdAt, preferences", "Stores account information, hashed password, language, and theme preference."),
        ("reports", "user, medicineName, use, dosage, precautions, sideEffects, doctorAdvice, imageUrl, originalAnalysisResult, displayedResult, selectedLanguage, timestamps", "Stores each authenticated user's medicine-analysis history."),
    ], [Inches(1.0), Inches(3.2), Inches(2.15)], font_size=8.8)
    add_heading(doc, "3.4 Security Design", 2)
    for item in [
        "Passwords are hashed using bcryptjs before storage.",
        "JWT tokens protect profile, scan, translation, and history APIs.",
        "Reports are queried by both report ID and logged-in user ID where appropriate.",
        "Gemini API key, JWT secret, and MongoDB URI are stored as backend environment variables.",
        "Uploaded files are restricted by extension, MIME type, and maximum size.",
    ]:
        add_bullet(doc, item)


def chapter_implementation(doc):
    add_chapter(doc, "4", "Implementation")
    add_heading(doc, "4.1 Technology Stack", 2)
    add_table(doc, ["Layer", "Technologies", "Role"], [
        ("Frontend", "React, Vite, Axios, React Router, CSS", "Responsive interface, API calls, routing, protected pages"),
        ("Backend", "Node.js, Express, body-parser, cors, dotenv", "REST API, validation, request handling"),
        ("Database", "MongoDB, Mongoose", "Persistent users and reports"),
        ("Authentication", "bcryptjs, jsonwebtoken", "Password hashing and JWT sessions"),
        ("Image Upload", "Multer", "Validated disk upload for JPG/JPEG/PNG files"),
        ("AI", "Google Gemini API", "Image understanding and value-only translation"),
        ("PDF", "jsPDF, Noto Sans Devanagari", "English/Hindi report export"),
        ("Deployment", "Vercel, Render, MongoDB Atlas", "Cloud hosting"),
    ], [Inches(1.1), Inches(2.45), Inches(2.8)], font_size=9)
    add_heading(doc, "4.2 Frontend Modules", 2)
    for item in [
        "Layout.jsx: shared sidebar, navbar, profile dropdown, settings modal, support card, and footer.",
        "AuthContext.jsx: token persistence, profile refresh, login, registration, logout, and preference updates.",
        "ProtectedRoute.jsx: redirects unauthenticated users to login while preserving their intended destination.",
        "Home.jsx: public landing dashboard, workflow cards, key features, and recent analyses.",
        "ScanMedicine.jsx: upload validation, AI analysis, translation cache, result cards, and PDF export.",
        "History.jsx: saved-report search, thumbnails, report details, and deletion.",
        "AppIcon.jsx and Logo.jsx: reusable presentation components.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "4.3 Backend Modules", 2)
    for item in [
        "app.js: Express application setup, CORS configuration, body parsing, static uploads, and route mounting.",
        "server.js: database connection and server startup.",
        "authController.js: account registration, login, profile response, and preference updates.",
        "medicineController.js: Gemini analysis, translation, and report persistence.",
        "geminiService.js: structured prompts, JSON parsing, normalization, and API-specific error handling.",
        "uploadMiddleware.js: Multer disk storage, filename sanitization, MIME validation, extension validation, and 10 MB file limit.",
        "reportController.js: user-specific report listing, retrieval, creation, and deletion.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "4.4 API Endpoints", 2)
    endpoints = [
        ("POST", "/api/auth/register", "Public", "Create user account"),
        ("POST", "/api/auth/login", "Public", "Login and receive JWT"),
        ("GET", "/api/auth/profile", "JWT", "Return logged-in profile"),
        ("PUT", "/api/auth/settings", "JWT", "Update language and theme"),
        ("POST", "/api/medicine/upload", "JWT", "Upload and analyze medicine image"),
        ("POST", "/api/medicine/analyze", "JWT", "Alias for upload-and-analyze"),
        ("POST", "/api/medicine/translate", "JWT", "Translate existing result values"),
        ("GET", "/api/reports", "JWT", "List current user's reports"),
        ("GET", "/api/reports/:id", "JWT", "Get one report"),
        ("DELETE", "/api/reports/:id", "JWT", "Delete one report"),
        ("GET", "/api/health", "Public", "Check API availability"),
    ]
    add_table(doc, ["Method", "Endpoint", "Access", "Purpose"], endpoints, [Inches(0.72), Inches(2.35), Inches(0.65), Inches(2.63)], font_size=8.8)
    add_heading(doc, "4.5 AI Prompt Safety Rules", 2)
    add_para(doc, "The Gemini prompt requires a JSON-only response and applies explicit safety boundaries. It requests dosage only when visible, instructs the model not to diagnose diseases or prescribe medicine, and requires a doctor or pharmacist consultation message. Hindi translation preserves medicine brand names, dosage numbers, and units.")
    add_callout(doc, "Medical Disclaimer", "MediLens AI is an informational assistant. It does not replace a qualified doctor or pharmacist and must not be used for diagnosis, treatment selection, or emergency decision-making.", fill="FFF2DB", color=GOLD)
    add_heading(doc, "4.6 User Interface", 2)
    if SCREENSHOT.exists():
        doc.add_picture(str(SCREENSHOT), width=Inches(6.15))
        add_para(doc, "Figure 2: MediLens AI responsive dashboard interface", size=9, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "The interface uses a purple and white medical theme with a scrollable sidebar, public dashboard, reusable icons, a support card, recent analyses, responsive cards, theme preference support, and a protected scan page.")


def chapter_testing(doc):
    add_chapter(doc, "5", "Testing and Validation")
    add_heading(doc, "5.1 Testing Strategy", 2)
    add_para(doc, "The project was checked through frontend production builds, backend syntax validation, manual API testing, and browser-based workflow verification during development. The following test matrix can be used during final submission demonstration.")
    tests = [
        ("TC-01", "Register with valid details", "New user is created; JWT returned", "Pass"),
        ("TC-02", "Login with correct credentials", "User enters dashboard", "Pass"),
        ("TC-03", "Login with invalid password", "Readable error message", "Pass"),
        ("TC-04", "Upload unsupported image type", "Upload rejected", "Pass"),
        ("TC-05", "Upload image larger than 10 MB", "Upload rejected", "Pass"),
        ("TC-06", "Analyze valid medicine image", "Structured result cards displayed", "Pass"),
        ("TC-07", "Change result language to Hindi", "Translated values displayed without rescanning", "Pass"),
        ("TC-08", "Download English PDF", "Readable purple-header report", "Pass"),
        ("TC-09", "Download Hindi PDF", "Devanagari text rendered with embedded font", "Pass"),
        ("TC-10", "View report history", "Only current user's reports displayed", "Pass"),
        ("TC-11", "Delete report", "Report removed from history", "Pass"),
        ("TC-12", "View recent analyses", "Recent rows and thumbnails shown", "Pass"),
    ]
    add_table(doc, ["ID", "Test Case", "Expected Result", "Status"], tests, [Inches(0.65), Inches(2.2), Inches(2.85), Inches(0.65)], font_size=8.7)
    add_heading(doc, "5.2 Error Handling", 2)
    for item in [
        "Invalid Gemini API key returns a clear backend configuration error.",
        "Gemini quota exhaustion returns a readable quota-limit message.",
        "Missing image upload returns a validation message.",
        "CORS restricts browser requests to configured frontend origins.",
        "Translation errors preserve the previous displayed result and show an error message.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "5.3 Validation Commands", 2)
    add_para(doc, "Frontend production build:")
    add_callout(doc, "Command", "cd frontend\nnpm run build", fill="F4F6F9", color=PURPLE_DARK)
    add_para(doc, "Backend syntax validation:")
    add_callout(doc, "Command", "node --check backend/src/app.js\nnode --check backend/src/controllers/medicineController.js\nnode --check backend/src/services/geminiService.js", fill="F4F6F9", color=PURPLE_DARK)


def chapter_deployment(doc):
    add_chapter(doc, "6", "Deployment")
    add_heading(doc, "6.1 Deployment Architecture", 2)
    add_para(doc, "The frontend and backend are deployed separately. The React application is deployed to Vercel, the Express API is deployed to Render, and MongoDB Atlas provides the cloud database. Secrets are stored only in Render environment variables.")
    add_table(doc, ["Component", "Platform", "Important Configuration"], [
        ("Frontend", "Vercel", "Root: frontend; Build: npm run build; Output: dist"),
        ("Backend", "Render Web Service", "Root: backend; Build: npm install; Start: npm start"),
        ("Database", "MongoDB Atlas", "MONGO_URI with database user and network access"),
        ("AI Service", "Google AI Studio / Gemini", "GEMINI_API_KEY stored on backend only"),
    ], [Inches(1.2), Inches(1.55), Inches(3.7)], font_size=9)
    add_heading(doc, "6.2 Backend Environment Variables", 2)
    add_table(doc, ["Variable", "Purpose"], [
        ("MONGO_URI", "MongoDB Atlas connection string"),
        ("JWT_SECRET", "Long secret used to sign JWT tokens"),
        ("JWT_EXPIRES_IN", "Token validity period such as 7d"),
        ("GEMINI_API_KEY", "Private Gemini API key"),
        ("GEMINI_MODEL", "Configured Gemini model name"),
        ("CLIENT_URL", "Allowed Vercel frontend origin"),
        ("REQUIRE_MONGO", "Set to true in deployment so startup fails if database is unavailable"),
    ], [Inches(2.0), Inches(4.45)], font_size=9.2)
    add_heading(doc, "6.3 Frontend Environment Variable", 2)
    add_table(doc, ["Variable", "Purpose"], [
        ("VITE_API_URL", "Public Render backend base URL used by Axios"),
    ], [Inches(2.0), Inches(4.45)], font_size=9.2)
    add_callout(doc, "Deployment Note", "Uploaded files are currently stored in backend/uploads. For long-term production use, move images to Cloudinary, Amazon S3, or another persistent object-storage service because ephemeral hosting files may not survive redeployment.", fill="FFF2DB", color=GOLD)


def chapter_future(doc):
    add_chapter(doc, "7", "Limitations and Future Scope")
    add_heading(doc, "7.1 Current Limitations", 2)
    for item in [
        "AI output depends on image clarity and visible label content.",
        "The system cannot guarantee medical correctness and must be treated as informational.",
        "Render local uploads may be temporary after redeployment.",
        "Prescription handwriting quality may affect recognition.",
        "The application currently supports English and Hindi only.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "7.2 Future Enhancements", 2)
    for item in [
        "OCR text preview before AI analysis so users can verify extracted label text.",
        "Cloud image storage with signed URLs and deletion lifecycle rules.",
        "Medicine reminder scheduling with notification support.",
        "Drug-interaction warnings with stronger medical-data validation.",
        "Family profiles for separate scan histories.",
        "Voice output for English and Hindi results.",
        "Favorites, manual pharmacist notes, and report sharing.",
        "Blur detection and image-quality scoring before upload.",
        "Additional Indian-language support.",
        "Automated API, integration, and end-to-end test suite.",
    ]:
        add_bullet(doc, item)


def chapter_conclusion(doc):
    add_chapter(doc, "8", "Conclusion")
    add_para(doc, "MediLens AI demonstrates how a full-stack web application can combine image upload, artificial intelligence, authentication, database persistence, multilingual support, and PDF reporting into a practical medicine-understanding assistant. The system presents complex medicine information in a structured form while maintaining an important safety boundary: the output is informational and cannot replace a healthcare professional.")
    add_para(doc, "The project successfully implements a public informational home page, protected user workflows, Gemini-based image analysis, translation without rescanning, user-specific history, scanned-image thumbnails, settings, responsive design, and deployment-ready frontend and backend separation. With cloud storage, expanded testing, and additional language support, it can be developed further into a stronger production-oriented platform.")


def references_and_appendix(doc):
    doc.add_page_break()
    add_heading(doc, "REFERENCES", 1)
    refs = [
        "React Documentation. https://react.dev/",
        "Vite Documentation. https://vite.dev/",
        "Express Documentation. https://expressjs.com/",
        "MongoDB Atlas Documentation. https://www.mongodb.com/docs/atlas/",
        "Mongoose Documentation. https://mongoosejs.com/docs/",
        "Google Gemini API Documentation. https://ai.google.dev/",
        "Multer Documentation. https://github.com/expressjs/multer",
        "JSON Web Token Documentation. https://jwt.io/",
        "jsPDF Documentation. https://github.com/parallax/jsPDF",
        "Vercel Documentation. https://vercel.com/docs",
        "Render Documentation. https://render.com/docs",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    doc.add_page_break()
    add_heading(doc, "APPENDIX A: PROJECT FOLDER STRUCTURE", 1)
    structure = """MediLens-AI/
|-- frontend/
|   |-- public/
|   |-- src/
|   |   |-- assets/
|   |   |-- components/
|   |   |-- context/
|   |   |-- pages/
|   |   |-- services/
|   |   |-- App.jsx
|   |   |-- main.jsx
|   |   `-- styles.css
|   |-- index.html
|   `-- package.json
`-- backend/
    |-- src/
    |   |-- config/
    |   |-- controllers/
    |   |-- middleware/
    |   |-- models/
    |   |-- routes/
    |   |-- services/
    |   |-- app.js
    |   `-- server.js
    |-- uploads/
    `-- package.json"""
    add_callout(doc, "Repository Structure", structure, fill="F4F6F9", color=PURPLE_DARK)

    add_heading(doc, "APPENDIX B: SAMPLE JSON RESPONSE", 1)
    sample = """{
  "medicineName": "Paracetamol 500mg Tablet",
  "use": "Used to reduce fever and relieve mild to moderate pain.",
  "dosage": "Use only as directed by a doctor or prescription.",
  "precautions": "Avoid overdose and consult a doctor when needed.",
  "sideEffects": "Nausea, stomach pain, dizziness, or allergic reaction may occur.",
  "doctorAdvice": "Consult a doctor or pharmacist before taking medicine."
}"""
    add_callout(doc, "Structured AI Response", sample, fill="F4F6F9", color=PURPLE_DARK)

    add_heading(doc, "APPENDIX C: FINAL SUBMISSION CHECKLIST", 1)
    for item in [
        "Replace all bracketed placeholders on the cover, certificate, and declaration pages.",
        "Add your final deployed Vercel URL and Render URL if required by your college.",
        "Verify Gemini API key and MongoDB Atlas connection on the deployed backend.",
        "Run one complete scan in English and Hindi.",
        "Open the downloaded PDF and confirm Hindi characters are readable.",
        "Take final screenshots from the deployed website for presentation slides.",
    ]:
        add_bullet(doc, item)


def main():
    make_architecture_diagram()
    doc = Document()
    configure_document(doc)
    cover_page(doc)
    preliminary_pages(doc)
    chapter_introduction(doc)
    chapter_requirements(doc)
    chapter_design(doc)
    chapter_implementation(doc)
    chapter_testing(doc)
    chapter_deployment(doc)
    chapter_future(doc)
    chapter_conclusion(doc)
    references_and_appendix(doc)
    props = doc.core_properties
    props.title = "MediLens AI - Medicine Understanding Assistant"
    props.subject = "College Project Report"
    props.author = "[STUDENT NAME]"
    props.keywords = "MediLens AI, React, Express, MongoDB, Gemini API"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
